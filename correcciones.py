import streamlit as st
import stripe
import os
from io import BytesIO
import requests
import jwt
import datetime
from textwrap import dedent
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# -------------------------------
# Función para crear el archivo DOCX
# -------------------------------
def create_docx(analysis, correction):
    """
    Crea un archivo DOCX con el análisis literario y la corrección de estilo.

    Args:
        analysis (str): Texto del análisis literario.
        correction (str): Texto de la corrección con justificaciones.

    Returns:
        BytesIO: Objeto BytesIO con el contenido del DOCX.
    """
    document = Document()

    # Título del Documento
    title = document.add_heading('Análisis Literario y Corrección de Estilo', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sección de Análisis Literario
    document.add_heading('📄 Análisis Literario', level=2)
    document.add_paragraph(analysis)

    # Sección de Corrección de Estilo
    document.add_heading('✍️ Corrección de Estilo, Ortográfica, Gramatical y de Puntuación con Justificaciones', level=2)
    
    # Procesar el texto de corrección para separar cambios y justificaciones
    pattern = re.compile(r'(.*?)\[(.*?)\]', re.DOTALL)
    matches = pattern.findall(correction)

    for original, justification in matches:
        para = document.add_paragraph()
        run_original = para.add_run(original)
        run_original.font.size = Pt(12)

        run_justification = para.add_run(f"[{justification}]")
        run_justification.font.size = Pt(12)
        run_justification.font.color.rgb = RGBColor(255, 0, 0)  # Rojo

    # Manejar cualquier texto que no coincida con el patrón
    non_matched_text = pattern.sub('', correction)
    if non_matched_text.strip():
        document.add_paragraph(non_matched_text)

    # Guardar el documento en un objeto BytesIO
    docx_io = BytesIO()
    document.save(docx_io)
    docx_io.seek(0)

    return docx_io

# -------------------------------
# Configuración de la página
# -------------------------------
st.set_page_config(
    page_title="Análisis Literario y Corrección de Estilo",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("🔍 Análisis Literario y Corrección de Estilo")

# Configurar la clave secreta de Stripe
stripe.api_key = st.secrets["STRIPE_SECRET_KEY"]

# ID del producto en Stripe
PRODUCT_ID = "prod_QwZPXT67PV3srt"  # Reemplaza con tu ID de producto real

# Obtener el secreto JWT de Streamlit secrets
JWT_SECRET = st.secrets["JWT_SECRET"]

# Función para generar un token JWT
def generate_jwt_token():
    payload = {
        "payment": "completed",
        "exp": datetime.datetime.utcnow() + datetime.timedelta(minutes=30)
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    return token

# Función para verificar el token JWT
def verify_jwt_token(token):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        if payload.get("payment") == "completed":
            return True
        return False
    except jwt.ExpiredSignatureError:
        st.error("El token ha expirado.")
        return False
    except jwt.InvalidTokenError:
        st.error("Token inválido.")
        return False

# Función para obtener el precio del producto
def get_price_for_product(product_id):
    try:
        prices = stripe.Price.list(product=product_id)
        if prices and prices['data']:
            return prices['data'][0].id
        else:
            st.error("No se encontraron precios para el producto.")
            return None
    except Exception as e:
        st.error(f"Error al recuperar el precio: {e}")
        return None

# Función para crear una sesión de pago en Stripe
def create_checkout_session(price_id):
    try:
        token = generate_jwt_token()
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price": price_id,
                "quantity": 1,
            }],
            mode="payment",
            success_url=f"{st.secrets['APP_URL']}?token={token}",
            cancel_url=f"{st.secrets['APP_URL']}?cancel=true",
        )
        return session
    except Exception as e:
        st.error(f"Error al crear la sesión de pago: {e}")
        return None

# -------------------------------
# Sidebar Instructions
# -------------------------------
st.sidebar.header("Instrucciones")
st.sidebar.markdown("""
**Esta aplicación te permite:**

- Realizar un análisis literario detallado de tu texto.
- Obtener recomendaciones de estilo específicas.
- Descargar una versión corregida de tu texto con justificaciones de los cambios realizados.

**Proceso:**

1. **Pagar por el servicio:** Una tarifa única de 9 USD.
2. **Subir tu texto:** Máximo 2000 palabras.
3. **Recibir análisis y corrección.**

**Aviso legal:**

Es responsabilidad del usuario verificar todos los cambios antes de utilizar el documento corregido.
""")

# -------------------------------
# Payment Handling
# -------------------------------

# Obtener parámetros de consulta
query_params = st.experimental_get_query_params()
token = query_params.get("token")
cancel = query_params.get("cancel")

if cancel:
    st.warning("El pago fue cancelado. Puedes intentarlo nuevamente.")

# Verificar si el usuario ya ha pagado
if token and verify_jwt_token(token[0]):
    st.success("¡Pago completado! Ahora puedes usar la aplicación.")
    acceso_concedido = True
else:
    acceso_concedido = False
    st.warning("Debes completar el pago antes de usar la aplicación.")

    if st.button("Pagar con Stripe"):
        price_id = get_price_for_product(PRODUCT_ID)
        if price_id:
            session = create_checkout_session(price_id)
            if session:
                st.markdown(f"[Proceder al pago]({session.url})", unsafe_allow_html=True)

# -------------------------------
# Literary Analysis and Correction
# -------------------------------

if acceso_concedido:
    st.header("📄 Análisis Literario y Corrección de Estilo")

    # Formulario de entrada
    with st.form(key='literary_analysis_form'):
        # Área de texto para el contenido
        text_input = st.text_area(
            "Pega tu texto (máximo 2000 palabras):",
            height=300,
            help="Asegúrate de que tu texto no exceda las 2000 palabras."
        )

        # Selección de género
        genre = st.selectbox(
            "Selecciona el género:",
            options=[
                "Fantasía", "Ciencia Ficción", "Misterio", "Romance",
                "Terror", "Aventura", "Drama", "Histórico", "Otro"
            ]
        )

        # Entrada de audiencia
        audience = st.text_input(
            "Define la audiencia:",
            help="Por ejemplo: adolescentes, adultos jóvenes, adultos, etc."
        )

        # Botón de envío
        submit_button = st.form_submit_button(label='Analizar y Corregir')

    # Función para contar palabras
    def count_words(text):
        return len(text.split())

    # Función para llamar a la API de Together para Análisis Literario
    def call_together_api_analysis(api_key, genre, audience, text):
        url = "https://api.together.xyz/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        # Construcción de los mensajes para la API con instrucciones claras y específicas
        messages = [
            {
                "role": "system",
                "content": dedent("""
                    Eres un crítico literario experto que proporciona análisis detallados y recomendaciones de estilo basadas en el género y la audiencia especificados.
                    **No debes corregir, modificar ni repetir el texto proporcionado.**
                    Tu única tarea es analizar el texto y ofrecer sugerencias de mejora enfocadas en aspectos literarios específicos como temas, desarrollo de personajes, estructura narrativa, tono y estilo.
                """)
            },
            {
                "role": "user",
                "content": dedent(f"""
                    Por favor, analiza el siguiente texto y proporciona una crítica literaria junto con recomendaciones de estilo específicas.

                    **Instrucciones adicionales:**
                    - No repitas el análisis previamente proporcionado.
                    - No corrijas ni modifiques el texto original de ninguna manera.
                    - Enfócate únicamente en proporcionar observaciones, críticas constructivas y sugerencias de mejora relacionadas directamente con el contenido del texto.
                    - Preserva todos los hipervínculos existentes en el texto. No agregues nuevos hipervínculos a menos que sean necesarios.
                    - No alteres las URLs de los hipervínculos existentes.
                    - Organiza el análisis en secciones claras como **Temas**, **Desarrollo de Personajes**, **Estructura Narrativa**, **Estilo y Tono**, etc.

                    **Género:** {genre}
                    **Audiencia:** {audience}

                    **Texto:**
                    {text}
                """)
            }
        ]

        payload = {
            "model": "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
            "messages": messages,
            "max_tokens": 2000,  # Ajusta según tus necesidades y límites de la API
            "temperature": 0.5,  # Reducida para respuestas más enfocadas
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["<|eot_id|>"],
            "stream": False  # Mantener como False para simplificar
        }

        try:
            response = requests.post(url, headers=headers, data=json.dumps(payload))
            response.raise_for_status()  # Esto lanzará una excepción si hay un error HTTP
            return response.json()
        except requests.exceptions.RequestException as e:
            st.error(f"Error al comunicarse con la API de Análisis: {e}")
            return None

    # Función para llamar a la API de Together para Corrección de Estilo y Ortografía con Justificaciones Inline
    def call_together_api_style_correction_with_justifications(api_key, analysis, text):
        url = "https://api.together.xyz/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        # Construcción de los mensajes para la API con instrucciones claras y específicas
        messages = [
            {
                "role": "system",
                "content": dedent("""
                    Eres un editor experto en corrección de estilo, ortografía, gramática y puntuación que revisa textos literarios.
                    **No debes realizar cambios que alteren el contenido original del autor.**
                    Tu tarea es corregir el estilo, ortografía, gramática y puntuación del texto proporcionado basado en el análisis y las recomendaciones previas.
                    **Preserva todos los hipervínculos existentes en el texto. No agregues nuevos hipervínculos a menos que sean necesarios. No alteres las URLs de los hipervínculos existentes.**
                    **Después de cada cambio realizado, añade una justificación entre corchetes y en color rojo.**
                """)
            },
            {
                "role": "user",
                "content": dedent(f"""
                    Basado en el siguiente análisis y recomendaciones, realiza una corrección de estilo del texto proporcionado. Incluye también correcciones ortográficas, gramaticales y de puntuación. Después de cada cambio realizado, añade una justificación entre corchetes y estilizada en color rojo.

                    **Análisis y Recomendaciones:**
                    {analysis}

                    **Texto Original:**
                    {text}

                    **Instrucciones adicionales:**
                    - No corrijas ni modifiques el contenido del texto.
                    - Enfócate únicamente en mejorar la claridad, el flujo, el estilo, la ortografía, la gramática y la puntuación.
                    - Preserva todos los hipervínculos existentes en el texto. No agregues nuevos hipervínculos a menos que sean necesarios.
                    - No alteres las URLs de los hipervínculos existentes.
                    - Para cada cambio realizado, proporciona una justificación detallada entre corchetes y estilizada en color rojo.
                    - Presenta el texto corregido con las justificaciones inline.
                """)
            }
        ]

        payload = {
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "messages": messages,
            "max_tokens": 3000,  # Aumentado para acomodar justificaciones
            "temperature": 0.5,  # Reducida para respuestas más enfocadas
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["<|eot_id|>"],
            "stream": False  # Mantener como False para simplificar
        }

        try:
            response = requests.post(url, headers=headers, data=json.dumps(payload))
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            st.error(f"Error al comunicarse con la API de Corrección de Estilo: {e}")
            return None

    # Acción al enviar el formulario
    if submit_button:
        # Validación de entrada
        if not text_input.strip():
            st.error("Por favor, pega tu texto para analizar y corregir.")
        elif not audience.strip():
            st.error("Por favor, define la audiencia.")
        else:
            word_count = count_words(text_input)
            if word_count > 2000:
                st.error(f"El texto excede el límite de 2000 palabras. Actualmente tiene {word_count} palabras.")
            else:
                # Mostrar spinner mientras se procesa la solicitud
                with st.spinner("Procesando tu solicitud..."):
                    # Obtener la API Key desde los secretos
                    try:
                        api_key = st.secrets["TOGETHER_API_KEY"]
                    except KeyError:
                        st.error("La clave de la API no está configurada correctamente en los secrets.")
                        st.stop()

                    # Primera llamada a la API para Análisis Literario
                    api_response_analysis = call_together_api_analysis(api_key, genre, audience, text_input)

                    if api_response_analysis:
                        # Extraer la respuesta del modelo para el análisis
                        try:
                            analysis = api_response_analysis['choices'][0]['message']['content']
                            st.subheader("📄 Análisis Literario")
                            st.write(analysis)
                        except (KeyError, IndexError):
                            st.error("Respuesta inesperada de la API de Análisis.")
                            analysis = None

                    # Segunda llamada a la API para Corrección de Estilo y Ortografía con Justificaciones Inline, si el análisis fue exitoso
                    if analysis:
                        api_response_correction = call_together_api_style_correction_with_justifications(api_key, analysis, text_input)

                        if api_response_correction:
                            # Extraer la respuesta del modelo para la corrección de estilo con justificaciones
                            try:
                                correction = api_response_correction['choices'][0]['message']['content']
                                st.subheader("✍️ Corrección de Estilo, Ortográfica, Gramatical y de Puntuación con Justificaciones")
                                # Renderizar el texto corregido con justificaciones en rojo
                                # Asegúrate de que la respuesta de la API esté formateada en Markdown con HTML para estilos
                                st.markdown(correction, unsafe_allow_html=True)
                            except (KeyError, IndexError):
                                st.error("Respuesta inesperada de la API de Corrección de Estilo.")

                            # Crear el archivo DOCX
                            docx_file = create_docx(analysis, correction)

                            # Proporcionar el botón de descarga
                            st.download_button(
                                label="📥 Descargar Análisis y Corrección (DOCX)",
                                data=docx_file,
                                file_name="Analisis_y_Correcion.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

    # -------------------------------
    # Optional: Reset Query Parameters
    # -------------------------------

    # Limpiar los parámetros de consulta después de procesar
    if acceso_concedido and (token or cancel):
        st.experimental_set_query_params()
